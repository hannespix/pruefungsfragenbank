import os
import sys
import socket
import re
import json
import requests
from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, send_file
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import db, Question, Exam, ExamItem, LLMConfig

# PyInstaller Trick: resource_path() Funktion
def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller"""
    try:
        # PyInstaller erstellt ein temporäres Verzeichnis und speichert den Pfad in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    
    return os.path.join(base_path, relative_path)


# Flask-App initialisieren - Templates/Static-Pfade für Dev und PyInstaller
if getattr(sys, 'frozen', False):
    # PyInstaller-Modus
    template_dir = resource_path('templates')
    static_dir = resource_path('static')
else:
    # Entwicklungsmodus
    template_dir = 'templates'
    static_dir = 'static'

app = Flask(__name__, template_folder=template_dir, static_folder=static_dir)
app.config['SECRET_KEY'] = 'your-secret-key-change-in-production'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(
    os.path.dirname(os.path.abspath(__file__)) if not getattr(sys, 'frozen', False) else os.path.dirname(sys.executable),
    'instance',
    'hortiexam.db'
)
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(
    os.path.dirname(os.path.abspath(__file__)) if not getattr(sys, 'frozen', False) else os.path.dirname(sys.executable),
    'instance',
    'uploads'
)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Erstelle Upload-Ordner falls nicht vorhanden
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(os.path.dirname(app.config['SQLALCHEMY_DATABASE_URI'].replace('sqlite:///', '')), exist_ok=True)

db.init_app(app)

# Erstelle Datenbank beim Start
with app.app_context():
    db.create_all()


def get_local_ip():
    """Ermittle die lokale IP-Adresse für LAN-Zugriff"""
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"


@app.route('/')
def index():
    """Hauptseite - Exam Builder"""
    exams = Exam.query.order_by(Exam.date_created.desc()).all()
    return render_template('index.html', exams=exams)


@app.route('/questions')
def questions():
    """API: Liste aller Fragen (filterbar)"""
    category = request.args.get('category', '')
    tag = request.args.get('tag', '')
    difficulty = request.args.get('difficulty', type=int)
    active_only = request.args.get('active_only', 'true') == 'true'
    
    query = Question.query
    
    if active_only:
        query = query.filter(Question.active == True)
    if category:
        query = query.filter(Question.category == category)
    if tag:
        query = query.filter(Question.tags.contains(tag))
    if difficulty:
        query = query.filter(Question.difficulty == difficulty)
    
    questions = query.order_by(Question.date_created.desc()).all()
    
    return jsonify([{
        'id': q.id,
        'content': q.content,
        'answer': q.answer,
        'category': q.category,
        'tags': q.tags.split(',') if q.tags else [],
        'difficulty': q.difficulty,
        'active': q.active
    } for q in questions])


@app.route('/exam/<int:exam_id>')
def exam_view(exam_id):
    """Ansicht einer Prüfung"""
    exam = Exam.query.get_or_404(exam_id)
    return render_template('exam.html', exam=exam)


@app.route('/exam/<int:exam_id>/items')
def exam_items(exam_id):
    """API: Items einer Prüfung"""
    exam = Exam.query.get_or_404(exam_id)
    items = ExamItem.query.filter_by(exam_id=exam_id).order_by(ExamItem.position).all()
    
    return jsonify([{
        'id': item.id,
        'content': item.snapshot_content,
        'answer': item.snapshot_answer,
        'points': item.points,
        'position': item.position,
        'original_question_id': item.original_question_id
    } for item in items])


@app.route('/exam/new', methods=['POST'])
def exam_new():
    """Neue Prüfung erstellen"""
    title = request.json.get('title', 'Neue Prüfung')
    exam = Exam(title=title, status='Draft')
    db.session.add(exam)
    db.session.commit()
    return jsonify({'id': exam.id, 'title': exam.title})


@app.route('/exam/<int:exam_id>/add_question', methods=['POST'])
def exam_add_question(exam_id):
    """Frage zur Prüfung hinzufügen (Snapshot-Pattern!)"""
    exam = Exam.query.get_or_404(exam_id)
    question_id = request.json.get('question_id')
    
    question = Question.query.get_or_404(question_id)
    
    # Snapshot erstellen: Content und Answer kopieren
    max_position = db.session.query(db.func.max(ExamItem.position)).filter_by(exam_id=exam_id).scalar() or -1
    
    exam_item = ExamItem(
        exam_id=exam_id,
        original_question_id=question_id,
        snapshot_content=question.content,  # SNAPSHOT!
        snapshot_answer=question.answer,    # SNAPSHOT!
        points=request.json.get('points', 1),
        position=max_position + 1
    )
    
    db.session.add(exam_item)
    db.session.commit()
    
    return jsonify({'success': True, 'item_id': exam_item.id})


@app.route('/exam/<int:exam_id>/remove_item/<int:item_id>', methods=['DELETE'])
def exam_remove_item(exam_id, item_id):
    """Item aus Prüfung entfernen"""
    item = ExamItem.query.filter_by(id=item_id, exam_id=exam_id).first_or_404()
    db.session.delete(item)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/exam/<int:exam_id>/reorder', methods=['POST'])
def exam_reorder(exam_id):
    """Reihenfolge der Items ändern"""
    item_ids = request.json.get('item_ids', [])
    for position, item_id in enumerate(item_ids):
        item = ExamItem.query.filter_by(id=item_id, exam_id=exam_id).first()
        if item:
            item.position = position
    db.session.commit()
    return jsonify({'success': True})


@app.route('/import', methods=['GET', 'POST'])
def import_questions():
    """Word-Dokument hochladen und Fragen importieren"""
    if request.method == 'GET':
        llm_configs = LLMConfig.query.filter_by(active=True).all()
        return render_template('import.html', llm_configs=llm_configs)
    
    if 'file' not in request.files:
        flash('Keine Datei ausgewählt', 'error')
        return redirect(url_for('import_questions'))
    
    file = request.files['file']
    if file.filename == '':
        flash('Keine Datei ausgewählt', 'error')
        return redirect(url_for('import_questions'))
    
    if file and file.filename.endswith('.docx'):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        try:
            use_llm = request.form.get('use_llm') == 'on'
            llm_config_id = request.form.get('llm_config_id', type=int)
            
            if use_llm and llm_config_id:
                # LLM-basierter Import
                count = import_from_word_with_llm(filepath, llm_config_id)
            else:
                # Klassischer Import
                count = import_from_word(filepath)
            
            flash(f'{count} Fragen erfolgreich importiert!', 'success')
        except Exception as e:
            flash(f'Fehler beim Import: {str(e)}', 'error')
        finally:
            if os.path.exists(filepath):
                os.remove(filepath)
    
    return redirect(url_for('import_questions'))


def import_from_word(filepath):
    """Word-Dokument einlesen und Fragen extrahieren"""
    doc = Document(filepath)
    count = 0
    current_question = None
    current_answer = None
    current_category = request.form.get('category', 'Allgemein')
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        # Ignoriere leere Zeilen
        if not text:
            continue
        
        # Frage erkennen (verschiedene Formate)
        if text.lower().startswith('frage:') or text.startswith('FRAGE:'):
            # Vorherige Frage speichern
            if current_question and current_answer:
                question = Question(
                    content=current_question.strip(),
                    answer=current_answer.strip(),
                    category=current_category,
                    active=True
                )
                db.session.add(question)
                count += 1
            
            # Neue Frage beginnen
            current_question = text.replace('Frage:', '').replace('FRAGE:', '').replace('frage:', '').strip()
            current_answer = None
        
        # Lösung erkennen (verschiedene Formate)
        elif text.lower().startswith('lösung:') or text.startswith('LÖSUNG:') or text.lower().startswith('loesung:'):
            if current_question:
                current_answer = text.replace('Lösung:', '').replace('LÖSUNG:', '').replace('lösung:', '').replace('Loesung:', '').strip()
            else:
                # Lösung ohne vorherige Frage - überspringen
                continue
        
        # Weiterer Text zur Frage oder Lösung
        elif current_question and not current_answer:
            # Weiterer Text zur Frage
            if current_question:
                current_question += '<br>' + text
        elif current_answer:
            # Weiterer Text zur Lösung
            current_answer += '<br>' + text
    
    # Letzte Frage speichern
    if current_question and current_answer:
        question = Question(
            content=current_question.strip(),
            answer=current_answer.strip(),
            category=current_category,
            active=True
        )
        db.session.add(question)
        count += 1
    
    db.session.commit()
    return count


def extract_text_from_word(filepath):
    """Extrahiert den gesamten Text aus einem Word-Dokument"""
    doc = Document(filepath)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())
    return '\n\n'.join(text_parts)


def call_llm_api(llm_config, text_content, category="Allgemein"):
    """Ruft die konfigurierte LLM-API auf und extrahiert Fragen"""
    try:
        headers = {
            'Content-Type': 'application/json'
        }
        
        # API-Key hinzufügen falls vorhanden
        if llm_config.api_key:
            if llm_config.provider == 'openai':
                headers['Authorization'] = f'Bearer {llm_config.api_key}'
            elif llm_config.provider == 'anthropic':
                headers['x-api-key'] = llm_config.api_key
                headers['anthropic-version'] = '2023-06-01'
            else:
                headers['Authorization'] = f'Bearer {llm_config.api_key}'
        
        # Zusätzliche Headers aus JSON parsen
        if llm_config.headers:
            try:
                extra_headers = json.loads(llm_config.headers)
                headers.update(extra_headers)
            except:
                pass
        
        # Prompt erstellen
        if llm_config.prompt_template:
            prompt = llm_config.prompt_template.replace('{text}', text_content).replace('{category}', category)
        else:
            prompt = f"""Analysiere folgenden Text und extrahiere alle Prüfungsfragen mit ihren Lösungen.

Text:
{text_content}

Bitte gib die Fragen und Lösungen im folgenden JSON-Format zurück:
{{
  "questions": [
    {{
      "content": "Die Frage hier",
      "answer": "Die Lösung hier",
      "category": "{category}",
      "tags": "Tag1, Tag2",
      "difficulty": 3
    }}
  ]
}}

Nur JSON zurückgeben, keine zusätzlichen Erklärungen."""

        # Request-Body je nach Provider
        if llm_config.provider == 'openai':
            body = {
                "model": llm_config.model or "gpt-4",
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.3
            }
        elif llm_config.provider == 'anthropic':
            body = {
                "model": llm_config.model or "claude-3-opus-20240229",
                "max_tokens": 4000,
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            }
        else:
            # Custom API - erwartet Standard-Format
            body = {
                "model": llm_config.model,
                "prompt": prompt,
                "temperature": 0.3,
                "max_tokens": 4000
            }
        
        # API-Call
        response = requests.post(
            llm_config.api_url,
            headers=headers,
            json=body,
            timeout=60
        )
        response.raise_for_status()
        
        # Response parsen
        data = response.json()
        
        # Response je nach Provider extrahieren
        if llm_config.provider == 'openai':
            content = data['choices'][0]['message']['content']
        elif llm_config.provider == 'anthropic':
            content = data['content'][0]['text']
        else:
            # Custom API - versuche verschiedene Formate
            content = data.get('response') or data.get('text') or data.get('content') or str(data)
        
        # JSON aus Response extrahieren (falls es in Markdown-Code-Blöcken ist)
        content = content.strip()
        if '```json' in content:
            content = content.split('```json')[1].split('```')[0].strip()
        elif '```' in content:
            content = content.split('```')[1].split('```')[0].strip()
        
        # JSON parsen
        result = json.loads(content)
        return result.get('questions', [])
        
    except Exception as e:
        raise Exception(f"LLM-API Fehler: {str(e)}")


def import_from_word_with_llm(filepath, llm_config_id):
    """Word-Dokument mit LLM analysieren und Fragen extrahieren"""
    llm_config = LLMConfig.query.get_or_404(llm_config_id)
    
    # Text aus Word extrahieren
    text_content = extract_text_from_word(filepath)
    
    if not text_content.strip():
        raise Exception("Das Word-Dokument enthält keinen Text")
    
    # LLM aufrufen
    category = request.form.get('category', 'Allgemein')
    questions_data = call_llm_api(llm_config, text_content, category)
    
    # Fragen in Datenbank speichern
    count = 0
    for q_data in questions_data:
        question = Question(
            content=q_data.get('content', '').strip(),
            answer=q_data.get('answer', '').strip(),
            category=q_data.get('category', category),
            tags=q_data.get('tags', ''),
            difficulty=q_data.get('difficulty', 3),
            active=True
        )
        if question.content and question.answer:
            db.session.add(question)
            count += 1
    
    db.session.commit()
    return count


@app.route('/settings', methods=['GET', 'POST'])
def settings():
    """Einstellungsseite für LLM-APIs"""
    if request.method == 'GET':
        configs = LLMConfig.query.order_by(LLMConfig.date_created.desc()).all()
        return render_template('settings.html', configs=configs)
    
    # POST: Neue Konfiguration speichern
    action = request.form.get('action')
    
    if action == 'create':
        config = LLMConfig(
            name=request.form.get('name'),
            api_url=request.form.get('api_url'),
            api_key=request.form.get('api_key', ''),
            model=request.form.get('model', ''),
            provider=request.form.get('provider', 'custom'),
            headers=request.form.get('headers', ''),
            prompt_template=request.form.get('prompt_template', ''),
            active=request.form.get('active') == 'on'
        )
        db.session.add(config)
        db.session.commit()
        flash('LLM-Konfiguration erfolgreich erstellt!', 'success')
    
    elif action == 'update':
        config_id = request.form.get('config_id', type=int)
        config = LLMConfig.query.get_or_404(config_id)
        config.name = request.form.get('name')
        config.api_url = request.form.get('api_url')
        config.api_key = request.form.get('api_key', '')
        config.model = request.form.get('model', '')
        config.provider = request.form.get('provider', 'custom')
        config.headers = request.form.get('headers', '')
        config.prompt_template = request.form.get('prompt_template', '')
        config.active = request.form.get('active') == 'on'
        db.session.commit()
        flash('LLM-Konfiguration erfolgreich aktualisiert!', 'success')
    
    elif action == 'delete':
        config_id = request.form.get('config_id', type=int)
        config = LLMConfig.query.get_or_404(config_id)
        db.session.delete(config)
        db.session.commit()
        flash('LLM-Konfiguration gelöscht!', 'success')
    
    return redirect(url_for('settings'))


@app.route('/settings/api/<int:config_id>')
def get_api_config(config_id):
    """API: Einzelne Konfiguration abrufen"""
    config = LLMConfig.query.get_or_404(config_id)
    return jsonify({
        'id': config.id,
        'name': config.name,
        'api_url': config.api_url,
        'api_key': config.api_key,
        'model': config.model,
        'provider': config.provider,
        'headers': config.headers,
        'prompt_template': config.prompt_template,
        'active': config.active
    })


@app.route('/export/<int:exam_id>')
def export_exam(exam_id):
    """Prüfung als Word-Dokument exportieren"""
    exam = Exam.query.get_or_404(exam_id)
    items = ExamItem.query.filter_by(exam_id=exam_id).order_by(ExamItem.position).all()
    
    # Word-Dokument erstellen
    doc = Document()
    
    # Kopfzeile
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = exam.title
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Titel
    title = doc.add_heading(exam.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datum
    date_para = doc.add_paragraph(f'Erstellt am: {exam.date_created.strftime("%d.%m.%Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Leerzeile
    
    # Fragen
    for idx, item in enumerate(items, 1):
        # Frage
        q_heading = doc.add_heading(f'Frage {idx} ({item.points} Punkte)', level=1)
        q_para = doc.add_paragraph()
        # HTML entfernen (einfache Version)
        content = item.snapshot_content
        content = content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        content = re.sub(r'<[^>]+>', '', content)  # Entferne alle HTML-Tags
        q_para.add_run(content)
        
        doc.add_paragraph()  # Leerzeile
    
    # Lösungen (neue Seite)
    doc.add_page_break()
    solutions_heading = doc.add_heading('Lösungen', 0)
    solutions_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for idx, item in enumerate(items, 1):
        sol_heading = doc.add_heading(f'Lösung {idx}', level=1)
        sol_para = doc.add_paragraph()
        answer = item.snapshot_answer
        answer = answer.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        answer = re.sub(r'<[^>]+>', '', answer)  # Entferne alle HTML-Tags
        sol_para.add_run(answer)
        doc.add_paragraph()
    
    # Speichern
    filename = f'exam_{exam_id}_{exam.title.replace(" ", "_")}.docx'
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc.save(filepath)
    
    return send_file(filepath, as_attachment=True, download_name=filename)


if __name__ == '__main__':
    local_ip = get_local_ip()
    port = 5000
    print(f"\n{'='*60}")
    print(f"HortiExam - Fragenbank für Gartenbau-Prüfungen")
    print(f"{'='*60}")
    print(f"Läuft auf:")
    print(f"  Lokal:    http://127.0.0.1:{port}")
    print(f"  LAN:      http://{local_ip}:{port}")
    print(f"{'='*60}\n")
    
    app.run(host='0.0.0.0', port=port, debug=False)
